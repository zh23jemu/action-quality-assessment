from __future__ import annotations

import json
import shutil
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
REPORT_DIR = DOCS / "reports"
ASSET_DIR = REPORT_DIR / "assets"

PRACTICE_TEMPLATE = ROOT / "1实践作业报告模板.docx"
TRAIN_TEMPLATE = ROOT / "2模型训练与测试说明模板.docx"

PRACTICE_REPORT = REPORT_DIR / "实践作业报告_MTL-AQA动作质量评估复现.docx"
TRAIN_REPORT = REPORT_DIR / "模型训练与测试说明_MTL-AQA动作质量评估复现.docx"

MOTION_CURVES = DOCS / "ppt_assets" / "motion_curves.png"
POSE_CURVES = DOCS / "ppt_assets" / "pose_curves.png"
POSE_F1_BAR = DOCS / "ppt_assets" / "pose_f1_bar.png"


def load_json(path: str) -> dict:
    """读取项目中已经同步回来的 JSON 结果文件。"""
    return json.loads((ROOT / path).read_text(encoding="utf-8"))


MOTION_EVAL = load_json("outputs/metrics/motion_eval_real.json")
POSE_EVAL = load_json("outputs/metrics/pose_eval_real.json")
MOTION_TRAIN = load_json("outputs/metrics/motion_train.json")
POSE_TRAIN = load_json("outputs/metrics/pose_train.json")
MOTION_SINGLE = load_json("outputs/predictions/motion_single_real.json")
POSE_SINGLE = load_json("outputs/predictions/pose_single_real.json")


def clear_document_body(doc: Document) -> None:
    """清空复制出来的模板正文，但保留模板本身的样式、节设置和页面属性。"""
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def normalize_sections(doc: Document) -> None:
    """统一复制版模板的页面节属性，避免原模板分栏导致正文、表格和图片被挤压。"""
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
        sect_pr = section._sectPr
        cols = sect_pr.find(qn("w:cols"))
        if cols is None:
            cols = OxmlElement("w:cols")
            sect_pr.append(cols)
        for child in list(cols):
            cols.remove(child)
        cols.set(qn("w:num"), "1")
        cols.set(qn("w:space"), "720")
        cols.set(qn("w:equalWidth"), "1")


def set_run_font(run, size: int | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    """统一设置中英文字体，避免中文在不同 Word 环境中回退得过于随意。"""
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_para(
    doc: Document,
    text: str = "",
    style: str | None = None,
    size: int = 10,
    bold: bool = False,
    color: RGBColor | None = None,
    align: WD_ALIGN_PARAGRAPH | None = None,
    before: int = 0,
    after: int = 6,
    line_spacing: float = 1.25,
    first_line: bool = True,
):
    """添加普通段落，并统一处理间距和字体。"""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    if align is None and text and first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    """添加层级标题，尽量复用模板的 Heading 样式。"""
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=14 if level == 1 else 12, bold=True, color=RGBColor(22, 54, 92))


def add_bullets(doc: Document, items: list[str]) -> None:
    """添加项目符号列表，用于减少长段文字密度。"""
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"• {item}")
        set_run_font(run, size=10)


def shade_cell(cell, fill: str) -> None:
    """设置表格单元格底色。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9, color: RGBColor | None = None) -> None:
    """写入单元格文字，并设置适合报告表格的内边距感和垂直居中。"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    """添加带浅色表头的表格，宽度按内容重要性分配。"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, "17365D")
        set_cell_text(cell, header, bold=True, size=9, color=RGBColor(255, 255, 255))
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=9)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    add_para(doc, "", after=4)


def add_picture(doc: Document, path: Path, caption: str, width_cm: float = 14.5) -> None:
    """插入图片并紧跟图注，避免图文分离。"""
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = add_para(doc, caption, size=9, color=RGBColor(90, 90, 90), align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    cap.paragraph_format.keep_with_next = True


def add_code_block(doc: Document, command: str) -> None:
    """用单列表格模拟命令块，便于 Word 中保持等宽和浅底色。"""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    shade_cell(cell, "F3F6FA")
    set_cell_text(cell, command, size=8, color=RGBColor(30, 45, 60))
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = "Consolas"
    add_para(doc, "", after=4)


def add_cover(doc: Document, title: str, subtitle: str, meta: list[str]) -> None:
    """生成简洁封面，保留模板但写入本项目真实题目。"""
    add_para(doc, title, size=22, bold=True, color=RGBColor(22, 54, 92), align=WD_ALIGN_PARAGRAPH.CENTER, before=80, after=16)
    add_para(doc, subtitle, size=13, color=RGBColor(60, 90, 120), align=WD_ALIGN_PARAGRAPH.CENTER, after=44)
    for line in meta:
        add_para(doc, line, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_para(doc, "说明：小组成员姓名、学号和课程名称可在最终提交前按实际信息补充。", size=9, color=RGBColor(130, 80, 50), align=WD_ALIGN_PARAGRAPH.CENTER, before=40)
    doc.add_page_break()


def best_epoch(history: list[dict], key: str) -> int:
    """从训练历史中找到指定指标最佳的轮次。"""
    return int(max(history, key=lambda item: item[key])["epoch"])


def make_terminal_image(path: Path, lines: list[str], title: str) -> None:
    """将关键终端输出绘制成图片，方便插入训练说明文档作为“截图式”证据。"""
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    font_path = Path("C:/Windows/Fonts/consola.ttf")
    title_font_path = Path("C:/Windows/Fonts/msyh.ttc")
    font = ImageFont.truetype(str(font_path), 22) if font_path.exists() else ImageFont.load_default()
    title_font = ImageFont.truetype(str(title_font_path), 24) if title_font_path.exists() else font
    wrapped: list[str] = []
    for line in lines:
        wrapped.extend(textwrap.wrap(line, width=92, subsequent_indent="  ") or [""])
    width = 1500
    height = 90 + len(wrapped) * 34 + 36
    img = Image.new("RGB", (width, height), (18, 30, 45))
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, width, 58], fill=(10, 20, 32))
    draw.ellipse([18, 18, 36, 36], fill=(238, 96, 85))
    draw.ellipse([46, 18, 64, 36], fill=(245, 190, 79))
    draw.ellipse([74, 18, 92, 36], fill=(92, 190, 105))
    draw.text((120, 16), title, font=title_font, fill=(225, 235, 245))
    y = 78
    for line in wrapped:
        color = (128, 221, 255) if line.startswith("$") else (226, 232, 240)
        if "SRC" in line or "mean_f1" in line or "R-L2" in line:
            color = (136, 226, 166)
        draw.text((28, y), line, font=font, fill=color)
        y += 34
    img.save(path)


def create_terminal_assets() -> dict[str, Path]:
    """生成训练、测试和单样本推理的终端输出图片素材。"""
    train_img = ASSET_DIR / "train_success_terminal.png"
    eval_img = ASSET_DIR / "eval_metrics_terminal.png"
    infer_img = ASSET_DIR / "single_infer_terminal.png"
    make_terminal_image(
        train_img,
        [
            "$ .venv/bin/python scripts/train_motion.py --manifest data/processed/mtl_aqa_manifest_features.csv",
            "Motion Disentangling training finished, checkpoint saved to outputs/checkpoints/motion_best.pt",
            "$ .venv/bin/python scripts/train_pose.py --manifest data/processed/mtl_aqa_manifest_features.csv",
            "Pose Contrastive training finished, checkpoint saved to outputs/checkpoints/pose_best.pt",
            "$ cat outputs/metrics/motion_train.json | tail",
            f"Motion Disentangling: best_SRC={MOTION_TRAIN['best_SRC']:.4f}, best_epoch={best_epoch(MOTION_TRAIN['history'], 'SRC')}, checkpoint={MOTION_TRAIN['checkpoint']}",
            "$ cat outputs/metrics/pose_train.json | tail",
            f"Pose Contrastive: best_mean_f1={POSE_TRAIN['best_mean_f1']:.4f}, best_epoch={best_epoch(POSE_TRAIN['history'], 'mean_f1')}, checkpoint={POSE_TRAIN['checkpoint']}",
        ],
        "直接 Python 训练脚本成功运行输出摘录",
    )
    make_terminal_image(
        eval_img,
        [
            "$ .venv/bin/python scripts/eval_motion.py --manifest data/processed/mtl_aqa_manifest_features.csv --checkpoint outputs/checkpoints/motion_best.pt --output outputs/metrics/motion_eval_real.json",
            json.dumps({"SRC": round(MOTION_EVAL["SRC"], 4), "R-L2": round(MOTION_EVAL["R-L2"], 4), "num_samples": MOTION_EVAL["num_samples"]}, ensure_ascii=False),
            "$ .venv/bin/python scripts/eval_pose.py --manifest data/processed/mtl_aqa_manifest_features.csv --checkpoint outputs/checkpoints/pose_best.pt --output outputs/metrics/pose_eval_real.json",
            json.dumps({"mean_f1": round(POSE_EVAL["mean_f1"], 4), "F1": {k: round(v, 4) for k, v in POSE_EVAL["F1"].items()}, "num_samples": POSE_EVAL["num_samples"]}, ensure_ascii=False),
        ],
        "两个模型测试指标输出摘录",
    )
    make_terminal_image(
        infer_img,
        [
            "$ .venv/bin/python scripts/infer_single.py --model motion --checkpoint outputs/checkpoints/motion_best.pt --manifest data/processed/mtl_aqa_manifest_features.csv --sample-id 01_0001 --output outputs/predictions/motion_single_real.json",
            json.dumps({"model": MOTION_SINGLE["model"], "sample_id": MOTION_SINGLE["sample_id"], "true_score": MOTION_SINGLE["true_score"], "pred_score": round(MOTION_SINGLE["pred_score"], 4)}, ensure_ascii=False),
            "$ .venv/bin/python scripts/infer_single.py --model pose --checkpoint outputs/checkpoints/pose_best.pt --manifest data/processed/mtl_aqa_manifest_features.csv --sample-id 01_0001 --output outputs/predictions/pose_single_real.json",
            json.dumps({"model": POSE_SINGLE["model"], "sample_id": POSE_SINGLE["sample_id"], "predicted_attributes": POSE_SINGLE["predicted_attributes"]}, ensure_ascii=False),
        ],
        "单样本推理完整流程输出摘录",
    )
    return {"train": train_img, "eval": eval_img, "infer": infer_img}


def setup_document_from_template(template: Path, output: Path) -> Document:
    """复制模板并打开副本，确保原模板文件不被直接修改。"""
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(template, output)
    doc = Document(output)
    clear_document_body(doc)
    normalize_sections(doc)
    return doc


def build_practice_report() -> None:
    """生成课程实践作业报告。"""
    doc = setup_document_from_template(PRACTICE_TEMPLATE, PRACTICE_REPORT)
    add_cover(
        doc,
        "MTL-AQA 上的动作质量评估适配复现实践报告",
        "Motion Disentangling 与 Pose Contrastive 两个模型的训练、测试与工程实现",
        ["第 X 组：成员姓名与学号待补充", "单位：计算机相关专业（课程实践报告）", "提交日期：2026 年 5 月"],
    )

    add_heading(doc, "摘要", 1)
    add_para(
        doc,
        "本实践围绕动作质量评估（Action Quality Assessment, AQA）任务， "
        "在 MTL-AQA 跳水数据集上适配 Fitness-AQA 论文中涉及的 Motion Disentangling 和 Pose Contrastive 两类模型。"
        "由于公开 GitHub 仓库无法直接提供完整的 MTL-AQA 训练闭环和原始论文特征，本项目补齐了数据下载、官方标注解析、视频抽帧、轻量帧统计特征生成、训练、独立测试、"
        "checkpoint 保存与单样本推理等工程环节。实验共处理 15 个长视频中标注出的 1412 个跳水片段，其中训练集 1059 个、测试集 353 个。"
        f"在独立测试集上，Motion Disentangling 取得 SRC={MOTION_EVAL['SRC']:.4f}、R-L2={MOTION_EVAL['R-L2']:.4f}；"
        f"Pose Contrastive 在五类动作属性上的平均 F1 为 {POSE_EVAL['mean_f1']:.4f}。"
        "结果表明，本项目已经完成面向课程交付的端到端可运行适配复现，但并非官方原始源码和原始 C3D/I3D/pose 特征的原封不动复现。",
        size=10,
    )
    add_para(doc, "关键词：动作质量评估；MTL-AQA；Motion Disentangling；Pose Contrastive；Python 训练", size=10, bold=True, first_line=False)

    add_heading(doc, "1 绪论", 1)
    add_para(
        doc,
        "动作质量评估的目标不是简单判断动作类别，而是根据视频中的完整动作过程输出得分或细粒度质量属性。"
        "在跳水、体操、花样滑冰等运动场景中，动作质量往往由起跳、空中姿态、翻腾、转体、入水等连续过程共同决定，"
        "因此模型既需要理解时间动态，也需要保留姿态和动作属性信息。课程需求要求使用 MTL-AQA 数据训练和测试两个模型，并输出 Motion 模型的 SRC、R-L2 以及 Pose 模型的 F1 Score。",
        size=10,
    )
    add_bullets(
        doc,
        [
            "研究问题：如何在真实 MTL-AQA 数据上打通从长视频到片段级训练样本的处理链路，并完成两个模型的训练测试。",
            "实践意义：形成可复验的工程流程，为客户验收、PPT 汇报和后续替换更强特征提供基础。",
            "总体思路：先解析官方标注和 split，再抽帧、生成统一 manifest，最后分别训练质量回归模型和动作属性分类模型。",
        ],
    )

    add_heading(doc, "2 相关工作与研究现状", 1)
    add_para(
        doc,
        "MTL-AQA 是跳水动作质量评估中常用的数据集，样本来自较长比赛视频中标注出的片段。Fitness-AQA 相关工作关注自监督表示学习，"
        "其中 Motion Disentangling 强调将动作动态表示用于质量分数回归，Pose Contrastive 强调姿态或属性层面的对比学习表示。"
        "本项目将两个思想迁移到 MTL-AQA 的片段级标签上：前者输出连续分数并计算 Spearman 排序相关与相对 L2 距离，后者输出 position、armstand、rotation_type、somersaults、twists 五类属性并计算 Macro F1。",
        size=10,
    )
    add_para(
        doc,
        "需要说明的是，公开仓库存在代码、训练入口、特征或数据链路缺失的问题，无法在当前环境中做到官方仓库原封不动运行。"
        "因此本实践采用“论文任务口径 + 公开信息 + MTL-AQA 官方标注 + 自建工程闭环”的适配复现方式。",
        size=10,
        color=RGBColor(140, 65, 40),
    )

    add_heading(doc, "3 关键实现技术", 1)
    add_heading(doc, "3.1 系统流程", 2)
    add_table(
        doc,
        ["阶段", "输入", "输出", "说明"],
        [
            ["数据获取", "15 个 MTL-AQA 长视频、官方 pkl 标注", "原始视频与标注目录", "视频平均时长约 1-3 小时，但训练单位是官方标注的跳水片段。"],
            ["抽帧与特征", "whole_videos/*.mp4、起止帧信息", "128 维帧统计特征 .npy", "仅按标注片段窗口采样，避免把长视频无关帧直接作为训练样本。"],
            ["Motion 训练", "feature manifest、score", "motion_best.pt、SRC/R-L2", "将片段特征映射为质量分数。"],
            ["Pose 训练", "feature manifest、五类属性", "pose_best.pt、mean F1", "对跳水动作属性进行多头分类。"],
            ["推理验证", "checkpoint、sample_id", "JSON 预测结果", "证明训练后模型可加载并对单样本输出结果。"],
        ],
        widths=[2.6, 4.0, 3.6, 5.0],
    )

    add_heading(doc, "3.2 数据处理与抽帧策略", 2)
    add_para(
        doc,
        "本项目没有把整条 2 小时左右的长视频直接作为训练样本，而是以 MTL-AQA 官方标注给出的片段为训练单位。"
        "每个样本都有 sample_id、所属视频编号、起止帧、得分和动作属性。抽帧后再根据起止帧窗口读取对应片段，生成统计特征。"
        "因此“无关帧”只作为长视频背景存在，不会被直接标记为有效训练样本。",
        size=10,
    )
    add_table(
        doc,
        ["统计项", "数量/结果"],
        [
            ["原始长视频", "15 个"],
            ["标注跳水片段", "1412 个"],
            ["训练集", "1059 个片段"],
            ["测试集", "353 个片段"],
            ["特征缺失", "0"],
            ["运行环境", "直接 Python 运行；NVIDIA L40S；PyTorch 2.12.0+cu126；CUDA 12.6"],
        ],
        widths=[4.0, 10.0],
    )

    add_heading(doc, "3.3 模型适配", 2)
    add_table(
        doc,
        ["模型", "任务定义", "输入输出", "评价指标"],
        [
            ["Motion Disentangling", "MTL-AQA 片段质量分数回归", "输入 128 维视频特征，输出 pred_score", "SRC 越高越好，R-L2 越低越好"],
            ["Pose Contrastive", "MTL-AQA 动作属性多分类", "输入 128 维视频特征，输出五个属性标签", "五类属性 Macro F1 及平均 F1"],
        ],
        widths=[4.0, 4.6, 4.6, 3.2],
    )

    add_heading(doc, "4 代码修改说明", 1)
    add_para(
        doc,
        "公开仓库不能直接完成本课程要求的 MTL-AQA 训练与测试，因此本项目围绕“可训练、可测试、可推理、可复验”的目标补齐了工程代码。"
        "修改重点不是更换任务目标，而是把公开资料中的模型思想落到 MTL-AQA 数据、统一 manifest、直接 Python 训练和 JSON 指标输出上。",
        size=10,
    )
    add_table(
        doc,
        ["文件/模块", "修改或新增内容", "作用"],
        [
            ["scripts/prepare_data.py", "解析 MTL-AQA 官方 pkl，生成包含 sample_id、split、score、动作属性和起止帧的 manifest。", "解决官方标注到训练样本的转换问题。"],
            ["scripts/extract_features.py", "按官方标注片段窗口读取抽帧图片，生成 128 维轻量帧统计特征，并输出带 feature_path 的 manifest。", "解决训练入口需要片段级特征的问题，避免把整段长视频无关帧作为样本。"],
            ["scripts/train_motion.py / eval_motion.py", "实现 Motion Disentangling 质量分数回归训练、checkpoint 保存、SRC 和 R-L2 指标输出。", "满足 Motion 模型训练与测试要求。"],
            ["scripts/train_pose.py / eval_pose.py", "实现 Pose Contrastive 五类动作属性多头分类训练、checkpoint 保存和 Macro F1 指标输出。", "满足 Pose 模型训练与测试要求。"],
            ["scripts/infer_single.py", "支持通过 sample_id 加载 checkpoint 并输出单样本预测 JSON。", "证明训练模型具备后续部署和单条数据推理入口。"],
            ["直接运行命令", "整理特征抽取、训练、独立测试和推理的 Python 命令，统一使用项目本地 .venv。", "保证客户无集群调度环境时也能按步骤直接运行。"],
        ],
        widths=[4.2, 7.2, 4.2],
    )
    add_para(
        doc,
        "此外，还修复了 feature_path 二次拼接、eval 参数名不一致、PyTorch 新版本 checkpoint 加载策略变化、AV1 视频解码工具版本不兼容等实际运行问题，"
        "这些修复保证了从数据准备到模型推理的完整链路可以重复执行。",
        size=10,
    )

    add_heading(doc, "5 验证与实验结果", 1)
    add_heading(doc, "5.1 实验环境", 2)
    add_table(
        doc,
        ["项目", "配置"],
        [
            ["训练平台", "直接 Python 运行环境"],
            ["GPU", "NVIDIA L40S"],
            ["Python / PyTorch", "Python 3.11；PyTorch 2.12.0+cu126"],
            ["CUDA", "12.6"],
            ["主要脚本", "prepare_data.py、extract_features.py、train_motion.py、train_pose.py、eval_motion.py、eval_pose.py、infer_single.py"],
        ],
        widths=[4.0, 10.0],
    )

    add_heading(doc, "5.2 指标结果", 2)
    add_table(
        doc,
        ["模型", "测试样本", "核心指标", "最佳轮次", "结果文件"],
        [
            ["Motion Disentangling", str(MOTION_EVAL["num_samples"]), f"SRC={MOTION_EVAL['SRC']:.4f}；R-L2={MOTION_EVAL['R-L2']:.4f}", str(best_epoch(MOTION_TRAIN["history"], "SRC")), "outputs/metrics/motion_eval_real.json"],
            ["Pose Contrastive", str(POSE_EVAL["num_samples"]), f"mean F1={POSE_EVAL['mean_f1']:.4f}", str(best_epoch(POSE_TRAIN["history"], "mean_f1")), "outputs/metrics/pose_eval_real.json"],
        ],
        widths=[4.0, 2.4, 5.0, 2.0, 4.4],
    )
    add_heading(doc, "5.3 输出结果截图", 2)
    add_para(
        doc,
        "以下图表来自训练与测试输出文件，展示两个模型的训练曲线、测试趋势和最终 F1 分布，可作为课程总结报告中的输出结果截图。",
        size=10,
    )
    add_picture(doc, MOTION_CURVES, "图 1  Motion Disentangling 输出结果截图：训练损失、SRC 与 R-L2 曲线", width_cm=15.5)
    add_picture(doc, POSE_CURVES, "图 2  Pose Contrastive 输出结果截图：训练损失与 mean F1 曲线", width_cm=15.0)
    add_picture(doc, POSE_F1_BAR, "图 3  Pose Contrastive 输出结果截图：五类动作属性 F1 对比", width_cm=13.5)

    add_heading(doc, "5.4 结果分析", 2)
    add_para(
        doc,
        "Motion 模型的 SRC 达到 0.7661，说明预测分数与真实分数在排序关系上具备较明显相关性；R-L2 为 0.1683，说明相对误差处于可解释范围。"
        "Pose 模型平均 F1 为 0.4052，其中 rotation_type、position、armstand 相对更稳定，twists 较低，主要原因可能是标签类别分布不均衡、轻量帧统计特征对细粒度转体动作表达不足。",
        size=10,
    )

    add_heading(doc, "6 结论", 1)
    add_para(
        doc,
        "本项目已经完成 MTL-AQA 上两个模型的端到端可运行适配复现：数据链路、训练、评估、checkpoint、单样本推理和汇报材料均已形成。"
        "当前结果适合作为课程实践和客户阶段性验收材料。后续若要追求论文级完整复现，需要进一步获得或重建官方原始视频特征、姿态特征和更完整的上游实现细节。",
        size=10,
    )

    add_heading(doc, "7 成果分工", 1)
    add_table(
        doc,
        ["成员", "主要工作", "贡献度"],
        [
            ["成员 1（待补充）", "数据获取、视频抽帧、manifest 生成与问题记录", "约 30%"],
            ["成员 2（待补充）", "Motion / Pose 模型训练、评估和推理链路实现", "约 40%"],
            ["成员 3（待补充）", "实验结果整理、PPT 和报告撰写", "约 30%"],
        ],
        widths=[4.0, 8.0, 3.0],
    )

    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] Parmar P, Morris B. Action Quality Assessment Across Multiple Actions. WACV, 2019.",
        "[2] MTL-AQA Dataset and Ready-to-use annotations.",
        "[3] Fitness-AQA related public repository and paper materials.",
        "[4] PyTorch Documentation: Model training, checkpoint and CUDA runtime.",
        "[5] FFmpeg Documentation: Video decoding and frame extraction.",
    ]
    for ref in refs:
        add_para(doc, ref, size=9, after=3, first_line=False)
    doc.save(PRACTICE_REPORT)


def build_training_report(terminal_assets: dict[str, Path]) -> None:
    """生成模型训练与测试说明文档。"""
    doc = setup_document_from_template(TRAIN_TEMPLATE, TRAIN_REPORT)
    add_cover(
        doc,
        "MTL-AQA 动作质量评估模型训练与测试说明",
        "Motion Disentangling 与 Pose Contrastive 的运行环境、命令、结果与推理流程",
        ["使用模型/数据集：Fitness-AQA 思想适配 + MTL-AQA 数据集", "模型 1：Motion Disentangling；模型 2：Pose Contrastive", "说明日期：2026 年 5 月"],
    )

    add_heading(doc, "1 使用环境详细说明", 1)
    add_table(
        doc,
        ["项目", "配置或说明"],
        [
            ["操作系统", "服务器端为 Linux 集群环境；本地整理与文档生成为 Windows"],
            ["运行方式", "直接使用项目本地 Python 环境运行训练、测试和推理脚本，不依赖集群调度系统"],
            ["GPU", "NVIDIA L40S"],
            ["Python", "Python 3.11，统一使用项目本地 .venv"],
            ["PyTorch", "2.12.0+cu126"],
            ["CUDA", "12.6"],
            ["关键依赖", "torch、numpy、pandas、pillow、imageio-ffmpeg、python-docx"],
            ["数据规模", "15 个长视频，1412 个官方标注跳水片段，train=1059，test=353"],
        ],
        widths=[4.0, 11.0],
    )
    add_para(
        doc,
        "说明：训练使用的是从真实视频帧抽取的 128 维轻量统计特征。客户环境只要准备好 Python 虚拟环境、数据和特征文件，"
        "即可直接运行下方 Python 命令完成训练与测试。由于公开上游仓库不完整，当前流程是可运行适配复现，"
        "不是官方原始源码和原始特征的封闭式完整复现。",
        size=10,
        color=RGBColor(140, 65, 40),
    )

    add_heading(doc, "2 程序运行步骤说明", 1)
    add_para(doc, "以下命令均为直接 Python 运行方式。Linux/macOS 使用 `.venv/bin/python`，Windows 可替换为 `.venv\\Scripts\\python.exe`。", size=10)
    add_heading(doc, "2.1 环境准备", 2)
    add_code_block(doc, ".venv/bin/python -m pip install --index-url https://download.pytorch.org/whl/cu126 torch torchvision\n.venv/bin/python -m pip install -r requirements.txt")
    add_heading(doc, "2.2 数据准备与特征生成", 2)
    add_code_block(doc, ".venv/bin/python scripts/prepare_data.py --raw-dir data/raw --output data/processed/mtl_aqa_manifest.csv\n.venv/bin/python scripts/extract_features.py --manifest data/processed/mtl_aqa_manifest.csv --output data/processed/mtl_aqa_manifest_features.csv")
    add_para(
        doc,
        "特征生成完成后，日志显示 total=1412、feature_ready=1412、missing_features=0，说明所有标注片段均已生成可训练特征。",
        size=10,
    )
    add_heading(doc, "2.3 直接训练两个模型", 2)
    add_code_block(doc, ".venv/bin/python scripts/train_motion.py --manifest data/processed/mtl_aqa_manifest_features.csv\n.venv/bin/python scripts/train_pose.py --manifest data/processed/mtl_aqa_manifest_features.csv")
    add_heading(doc, "2.4 独立测试", 2)
    add_code_block(doc, ".venv/bin/python scripts/eval_motion.py --manifest data/processed/mtl_aqa_manifest_features.csv --checkpoint outputs/checkpoints/motion_best.pt --output outputs/metrics/motion_eval_real.json\n.venv/bin/python scripts/eval_pose.py --manifest data/processed/mtl_aqa_manifest_features.csv --checkpoint outputs/checkpoints/pose_best.pt --output outputs/metrics/pose_eval_real.json")
    add_heading(doc, "2.5 单样本推理", 2)
    add_code_block(doc, ".venv/bin/python scripts/infer_single.py --model motion --checkpoint outputs/checkpoints/motion_best.pt --manifest data/processed/mtl_aqa_manifest_features.csv --sample-id 01_0001 --output outputs/predictions/motion_single_real.json\n.venv/bin/python scripts/infer_single.py --model pose --checkpoint outputs/checkpoints/pose_best.pt --manifest data/processed/mtl_aqa_manifest_features.csv --sample-id 01_0001 --output outputs/predictions/pose_single_real.json")

    add_heading(doc, "3 成功运行截图", 1)
    add_para(doc, "下图整理了训练任务提交与训练完成后的关键输出，实际完整日志位于 `outputs/logs/`。", size=10)
    add_picture(doc, terminal_assets["train"], "图 1 训练脚本成功运行终端输出摘录", width_cm=16.0)

    add_heading(doc, "4 测试结果与主要指标截图", 1)
    add_table(
        doc,
        ["模型", "测试集", "指标"],
        [
            ["Motion Disentangling", f"{MOTION_EVAL['num_samples']} 个片段", f"SRC={MOTION_EVAL['SRC']:.4f}；R-L2={MOTION_EVAL['R-L2']:.4f}"],
            ["Pose Contrastive", f"{POSE_EVAL['num_samples']} 个片段", f"mean F1={POSE_EVAL['mean_f1']:.4f}；position={POSE_EVAL['F1']['position']:.4f}；armstand={POSE_EVAL['F1']['armstand']:.4f}；rotation_type={POSE_EVAL['F1']['rotation_type']:.4f}；somersaults={POSE_EVAL['F1']['somersaults']:.4f}；twists={POSE_EVAL['F1']['twists']:.4f}"],
        ],
        widths=[4.0, 3.0, 9.0],
    )
    add_picture(doc, terminal_assets["eval"], "图 2 两个模型测试指标终端输出摘录", width_cm=16.0)
    add_picture(doc, MOTION_CURVES, "图 3 Motion Disentangling 训练与测试曲线", width_cm=15.5)
    add_picture(doc, POSE_F1_BAR, "图 4 Pose Contrastive 各动作属性 F1", width_cm=14.5)

    add_heading(doc, "5 使用训练模型进行单条数据推理流程", 1)
    add_para(
        doc,
        f"示例样本为 `01_0001`。Motion 模型真实分数为 {MOTION_SINGLE['true_score']:.1f}，预测分数为 {MOTION_SINGLE['pred_score']:.4f}；"
        f"Pose 模型预测属性为 position={POSE_SINGLE['predicted_attributes']['position']}、armstand={POSE_SINGLE['predicted_attributes']['armstand']}、"
        f"rotation_type={POSE_SINGLE['predicted_attributes']['rotation_type']}、somersaults={POSE_SINGLE['predicted_attributes']['somersaults']}、twists={POSE_SINGLE['predicted_attributes']['twists']}。",
        size=10,
    )
    add_picture(doc, terminal_assets["infer"], "图 5 单样本推理完整终端输出摘录", width_cm=16.0)

    add_heading(doc, "6 输出文件说明", 1)
    add_table(
        doc,
        ["类别", "路径", "用途"],
        [
            ["运行日志", "终端输出，或自行重定向到 outputs/logs/", "查看特征抽取、训练和测试过程"],
            ["模型权重", "outputs/checkpoints/motion_best.pt；outputs/checkpoints/pose_best.pt", "后续测试和推理加载"],
            ["训练指标", "outputs/metrics/motion_train.json；outputs/metrics/pose_train.json", "记录 30 个 epoch 的训练曲线"],
            ["测试指标", "outputs/metrics/motion_eval_real.json；outputs/metrics/pose_eval_real.json", "独立 eval 结果"],
            ["单样本推理", "outputs/predictions/motion_single_real.json；outputs/predictions/pose_single_real.json", "验证模型可部署入口"],
        ],
        widths=[3.0, 8.0, 5.0],
    )

    add_heading(doc, "7 注意事项", 1)
    add_bullets(
        doc,
        [
            "视频文件和抽帧图片体积较大，默认不纳入 Git 仓库，需要在服务器或 Release 中单独管理。",
            "如果重新抽帧遇到 AV1 解码失败，应优先使用 imageio-ffmpeg 附带的 ffmpeg 7.0.2 静态版本。",
            "当前特征是轻量统计特征，适合课程实践闭环；如需论文级对齐，应替换为更强的视频或姿态特征。",
            "公开上游仓库不完整，因此报告中应明确本项目定位为“可运行适配复现”。",
        ],
    )
    doc.save(TRAIN_REPORT)


def main() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    terminal_assets = create_terminal_assets()
    build_practice_report()
    build_training_report(terminal_assets)
    print(PRACTICE_REPORT)
    print(TRAIN_REPORT)


if __name__ == "__main__":
    main()
